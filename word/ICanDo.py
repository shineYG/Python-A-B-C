import docx
from docx.shared import Cm
from docx.shared import Pt
from PIL import Image, ImageDraw
from io import BytesIO
from docx.shared import RGBColor

recordset = [{'Qty':3,'Name':'Fish','Desc':'Tom'},
             {'Qty':8,'Name':'Cheese','Desc':'Jerry'},
             {'Qty':5,'Name':'Bacon','Desc':'Garfield'}]

document = docx.Document()

# title and body
document.add_heading('Document Title', 0)
p = document.add_paragraph('A plain paragraph having some')
p.add_run(' bold').bold = True
p.add_run('and some ')
p.add_run('italic.').italic = True
document.add_heading('Heading, level 1')
document.add_paragraph('Intense quote', style ='Intense Quote')

# list
document.add_paragraph('first item in unordered list', style = 'List Bullet')
document.add_paragraph('first item in ordered list', style = 'List Number')

# image
document.add_picture('classAndObject.png', width = Cm(10))
document.add_page_break()

# table
table = document.add_table(rows = 1, cols = 3, style = 'Table Grid')

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Name'
hdr_cells[2].text = 'Desc'

for item in recordset:
	row_cells = table.add_row().cells
	row_cells[0].text = str(item['Qty'])
	row_cells[1].text = item['Name']
	row_cells[2].text = item['Desc']


document.add_page_break()

# draw circle
p = document.add_paragraph()
r = p.add_run()
img_size = 20
for x in range(255):
	im = Image.new('RGB', (img_size, img_size), 'white')
	draw_obj = ImageDraw.Draw(im)
	draw_obj.ellipse((0,0,img_size-1,img_size-1), fill=255-x)
	# save image to buffer
	fake_buf_file = BytesIO()
	im.save(fake_buf_file, "png")
	r.add_picture(fake_buf_file)
	fake_buf_file.close()

document.add_page_break()

# font with changing color
p = document.add_paragraph()
text = '一个人的命运当然要靠自我奋斗，但是也要考虑到历史的进程。'
for i, ch in enumerate(text):
	run = p.add_run(ch)
	font  = run.font
	font.name = 'Silom'
	font.size = Pt(16)
	font.color.rgb = RGBColor(i*10%200 + 55, i*20%200 + 55, i*30%200 + 55)


document.save('ICanDo.docx')